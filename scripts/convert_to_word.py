"""
convert_to_word.py — VUCA AEB Extension
=========================================
Assembles all paper draft sections into a single AJAE-formatted Word document.

Section order:
  Abstract → Introduction → Literature Review → Methods →
  Results → Discussion → Conclusion → References → Tables

Formatting:
  - 12pt Times New Roman body, double-spaced, 1-inch margins
  - Section headings: 12pt bold
  - Subsection headings: 12pt bold italic
  - Math equations: converted to readable plain-text with [EQ n] markers
    (replace with proper Word equations before submission)
  - [CHECK] and [CITE] markers preserved for manual resolution
  - "Notes for revision" and draft metadata headers stripped

Output:
  output/VUCA_paper_draft.docx

Usage:
  python scripts/convert_to_word.py
"""

import re
import pathlib
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(pathlib.Path(__file__).parent))
import config

DRAFTS = config.PROJECT_ROOT / "quality_reports" / "drafts"
OUTPUT_DOCX = config.OUTPUT_DIR / "VUCA_paper_draft.docx"

# ── Equation counter ─────────────────────────────────────────────────────────
EQ_COUNTER = [0]


def clean_inline_math(text: str) -> str:
    """Replace inline $...$ with plain-text approximation."""
    def repl(m):
        inner = m.group(1)
        inner = re.sub(r"\\text\{([^}]+)\}", r"\1", inner)
        inner = re.sub(r"\\overline\{([^}]+)\}", r"\1̄", inner)
        inner = re.sub(r"\\bar\{([^}]+)\}", r"\1̄", inner)
        inner = re.sub(r"\\[a-zA-Z]+", "", inner)
        inner = re.sub(r"[{}]", "", inner)
        return inner.strip()
    return re.sub(r"\$([^$]+)\$", repl, text)


# ── Document formatting helpers ───────────────────────────────────────────────

def set_margins(doc: Document, inches: float = 1.0) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def set_font(run, size_pt: int = 12, bold: bool = False,
             italic: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph(doc: Document, text: str = "", style: str = "Normal",
                  bold: bool = False, italic: bool = False,
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  size_pt: int = 12, space_after_pt: int = 0) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.space_before = Pt(0)
    # Double spacing
    p.paragraph_format.line_spacing = Pt(24)  # 2 × 12pt

    if text:
        _add_inline_text(p, text, bold=bold, italic=italic, size_pt=size_pt)
    return p


def _add_inline_text(para, text: str, bold: bool = False,
                     italic: bool = False, size_pt: int = 12) -> None:
    """Add text to paragraph, handling **bold** and *italic* markers."""
    # Split on bold (**...**) and italic (*...*)
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|~~[^~]+~~)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_font(run, size_pt=size_pt, bold=True, italic=italic)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            set_font(run, size_pt=size_pt, bold=bold, italic=True)
        elif part.startswith("~~") and part.endswith("~~"):
            run = para.add_run(part[2:-2])
            set_font(run, size_pt=size_pt, bold=bold, italic=italic)
            run.font.strike = True
        else:
            run = para.add_run(part)
            set_font(run, size_pt=size_pt, bold=bold, italic=italic)


def add_heading(doc: Document, text: str, level: int) -> None:
    """Add section/subsection/subsubsection heading."""
    if level == 1:
        p = add_paragraph(doc, text, bold=True, size_pt=14,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    elif level == 2:
        p = add_paragraph(doc, text, bold=True, size_pt=12, space_after_pt=2)
    elif level == 3:
        p = add_paragraph(doc, text, bold=True, italic=True, size_pt=12, space_after_pt=2)
    else:
        p = add_paragraph(doc, text, italic=True, size_pt=12, space_after_pt=2)


def add_equation(doc: Document, latex: str) -> None:
    """
    Add a display equation as LaTeX source in Courier New.
    Preserves the raw LaTeX so it can be:
      - Copy-pasted into a LaTeX template, or
      - Converted in Word via Alt+= (Word's LaTeX equation input)
    """
    EQ_COUNTER[0] += 1
    src = latex.strip().strip("$$").strip()

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Equation number label
    run_num = p.add_run(f"({EQ_COUNTER[0]})  ")
    run_num.font.name = "Times New Roman"
    run_num.font.size = Pt(11)

    # LaTeX source in monospace
    run_eq = p.add_run(src)
    run_eq.font.name = "Courier New"
    run_eq.font.size = Pt(10)

    # Tip annotation on first equation only
    if EQ_COUNTER[0] == 1:
        tip = doc.add_paragraph()
        tip.paragraph_format.left_indent = Inches(0.5)
        tip.paragraph_format.space_after = Pt(2)
        tip.paragraph_format.line_spacing = Pt(14)
        run_tip = tip.add_run(
            "[To render: In Word, place cursor after equation, press Alt+=, "
            "then paste the LaTeX source into the equation editor]"
        )
        run_tip.font.name = "Times New Roman"
        run_tip.font.size = Pt(9)
        run_tip.font.italic = True
        run_tip.font.color.rgb = None  # default grey not supported; leave as black


def add_table_from_markdown(doc: Document, rows: list[str]) -> None:
    """Convert markdown table rows to a Word table."""
    # Parse rows
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    # Remove separator row (---|---|---)
    parsed = [r for r in parsed if not all(re.match(r"^[-:]+$", c) for c in r)]
    if not parsed:
        return

    n_cols = max(len(r) for r in parsed)
    # Pad short rows
    parsed = [r + [""] * (n_cols - len(r)) for r in parsed]

    tbl = doc.add_table(rows=len(parsed), cols=n_cols)
    tbl.style = "Table Grid"

    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row):
            cell = tbl.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = Pt(18)
            clean = clean_inline_math(cell_text)
            _add_inline_text(p, clean, bold=(i == 0), size_pt=10)

    doc.add_paragraph()  # spacing after table


# ── Markdown section parser ───────────────────────────────────────────────────

def parse_and_add(doc: Document, md_text: str,
                  skip_draft_header: bool = True,
                  skip_notes: bool = True) -> None:
    """
    Parse markdown and add content to doc.

    Skips:
      - Draft metadata header lines (Status, Target, Note lines at top)
      - "Notes for revision" / "Notes for Revision" sections
      - Section end markers (*[END OF...*)
      - Horizontal rules (---)
    """
    lines = md_text.splitlines()

    in_notes_section = False
    in_table = False
    table_rows: list[str] = []
    equation_buffer: list[str] = []
    in_equation = False
    in_code_block = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Skip code blocks ─────────────────────────────────────
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # ── Stop at Notes for revision ───────────────────────────
        if re.match(r"^##+ Notes for revision", stripped, re.IGNORECASE):
            in_notes_section = True
        if in_notes_section:
            i += 1
            continue

        # ── Skip draft metadata lines ─────────────────────────────
        if skip_draft_header and re.match(
            r"^\*\*(Target|Status|Note|Word count)\*\*:", stripped
        ):
            i += 1
            continue

        # ── Skip end-of-section markers ───────────────────────────
        if stripped.startswith("*[END OF") or stripped.startswith("[END OF"):
            i += 1
            continue

        # ── Skip horizontal rules ─────────────────────────────────
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # ── Display equations $$...$$ ─────────────────────────────
        if stripped == "$$" or stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            # Single-line $$...$$ or block start
            if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
                add_equation(doc, stripped)
                i += 1
                continue
            # Multi-line: toggle
            if not in_equation:
                in_equation = True
                equation_buffer = []
            else:
                in_equation = False
                add_equation(doc, " ".join(equation_buffer))
                equation_buffer = []
            i += 1
            continue

        if in_equation:
            equation_buffer.append(stripped)
            i += 1
            continue

        # ── Tables ────────────────────────────────────────────────
        if stripped.startswith("|"):
            table_rows.append(stripped)
            i += 1
            continue
        else:
            if table_rows:
                add_table_from_markdown(doc, table_rows)
                table_rows = []

        # ── Headings ──────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_inline_math(heading_match.group(2))
            # Skip the top-level file title (it's the section title, not paper title)
            if level == 1 and skip_draft_header:
                i += 1
                continue
            add_heading(doc, text, level)
            i += 1
            continue

        # ── Bullet lists ──────────────────────────────────────────
        if re.match(r"^[-*+]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^[-*+]\s+", "", stripped)
            text = re.sub(r"^\d+\.\s+", "", stripped)
            text = clean_inline_math(text)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(0)
            _add_inline_text(p, text, size_pt=12)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────
        text = clean_inline_math(stripped)
        add_paragraph(doc, text, space_after_pt=0)
        i += 1

    # Flush trailing table
    if table_rows:
        add_table_from_markdown(doc, table_rows)


# ── Reference list ────────────────────────────────────────────────────────────

REFERENCES = """Baker, S.R., Bloom, N. and Davis, S.J. (2016). Measuring Economic Policy Uncertainty. Quarterly Journal of Economics 131(4): 1593–1636.

Baran, B.E. and Woznyj, H.M. (2020). Managing VUCA: The Human Dynamics of Agility. Organizational Dynamics 49(2): 100787.

Bennett, N. and Lemoine, G.J. (2014). What a Difference a Word Makes: Understanding Threats to Performance in a VUCA World. Business Horizons 57(3): 311–317.

Bennis, W.G. and Nanus, B. (1985). Leaders: The Strategies for Taking Charge. Harper & Row.

Birru, J. and Young, T. (2023). The Real Effects of Sentiment and Uncertainty. Working Paper, Fisher College of Business, Ohio State University.

Blei, D.M., Ng, A.Y. and Jordan, M.I. (2003). Latent Dirichlet Allocation. Journal of Machine Learning Research 3: 993–1022.

Curtin, R. (2007). Consumer Sentiment Surveys: Worldwide Review and Assessment. Journal of Business Cycle Measurement and Analysis 2007(1). [CHECK: confirm exact reference]

Dixit, A.K. and Pindyck, R.S. (1994). Investment under Uncertainty. Princeton University Press.

Ellsberg, D. (1961). Risk, Ambiguity, and the Savage Axioms. Quarterly Journal of Economics 75(4): 643–669.

Featherstone, A.M. and Baker, T.G. (1987). An Examination of Farm Sector Real Asset Dynamics. American Journal of Agricultural Economics 69(3): 532–546.

Fridgeirsson, T.V., Kristjansdottir, B.H. and Ingason, H.T. (2021). An Alternative Risk Assessment Routine for Decision Making; Towards a VUCA Meter. In R. Cuevas, C.-N. Bodea, & P. Torres-Lima (Eds.), Research on Project, Programme and Portfolio Management (pp. 41–54). Springer.

Gentzkow, M., Kelly, B. and Taddy, M. (2019). Text as Data. Journal of Economic Literature 57(3): 535–574.

Goss, E. and Natvig, B. (2023). Creighton University Rural Mainstreet Economy Survey. [CHECK: confirm exact citation format]

Granger, C.W.J. (1969). Investigating Causal Relations by Econometric Models and Cross-Spectral Methods. Econometrica 37(3): 424–438.

Hartigan, J.A. and Hartigan, P.M. (1985). The Dip Test of Unimodality. Annals of Statistics 13(1): 70–84.

Knight, F.H. (1921). Risk, Uncertainty and Profit. Houghton Mifflin.

Leduc, S. and Liu, Z. (2016). Uncertainty Shocks Are Aggregate Demand Shocks. Journal of Monetary Economics 82: 20–35. [CHECK: confirm this is the right Leduc/Liu paper]

Lippsmeyer, M., Langemeier, M., Mintert, J. and Thompson, N. (2024). Factors Influencing Producer Sentiment. Journal of ASFMRA 2024: 58–70.

Loughran, T. and McDonald, B. (2011). When Is a Liability Not a Liability? Textual Analysis, Dictionaries, and 10-Ks. Journal of Finance 66(1): 35–65.

Mintert, J. and Widmar, D.A. (2016). Purdue University/CME Group Ag Economy Barometer. [CHECK: confirm exact working paper/report citation]

Moss, C.B. and Katchova, A.L. (2005). Farmland Valuation and Asset Performance. Agricultural Finance Review 65(1): 119–130. [CHECK: confirm exact year/journal]

Nowzohour, L. and Stracca, L. (2020). More Than a Feeling: Confidence, Uncertainty, and Macroeconomic Fluctuations. Journal of Economic Surveys 34(4): 1–36.

Patnaik, S. (2020). Applied Machine Learning and Management of Volatility, Uncertainty, Complexity & Ambiguity (V.U.C.A). Journal of Intelligent & Fuzzy Systems. https://journals.sagepub.com/doi/full/10.3233/JIFS-179915

Rzepczynski, M.S. (2025). Clarifying the Assessment of Risk: VUCA (Volatility, Uncertainty, Complexity, and Ambiguity). SSRN Working Paper No. 5217110.

Sims, C.A. (1980). Macroeconomics and Reality. Econometrica 48(1): 1–48.

[CHECK: Add — "Being Affected By VUCA Factors? Developing The Perceived VUCA Exposure Scale" — find full citation]
[CHECK: Add — "Clarifying the Conceptual Map of VUCA: A Systematic Review" — find full citation]"""


# ── Main assembly ─────────────────────────────────────────────────────────────

def read_draft(filename: str) -> str:
    path = DRAFTS / filename
    if not path.exists():
        print(f"  [WARN] Not found: {path}")
        return ""
    return path.read_text(encoding="utf-8")


def main():
    print("=" * 60)
    print("convert_to_word.py — Assembling paper draft")
    print("=" * 60)

    doc = Document()
    set_margins(doc, inches=1.0)

    # ── Default style ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Title page ────────────────────────────────────────────────
    p = add_paragraph(
        doc,
        "Measuring VUCA in U.S. Agriculture: Decomposing Agricultural "
        "Producer Sentiment into Volatility, Uncertainty, Complexity, "
        "and Ambiguity",
        bold=True, size_pt=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=12,
    )

    add_paragraph(doc, "[Author names — blinded for review]",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)
    add_paragraph(doc, "Purdue University",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)
    add_paragraph(doc, "",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)
    add_paragraph(doc, "Draft — March 2026. Do not cite without permission.",
                  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11)

    doc.add_page_break()

    # ── Abstract ──────────────────────────────────────────────────
    print("  Adding: Abstract")
    abstract_md = read_draft("conclusion_abstract_draft.md")
    # Extract just the abstract block
    abs_match = re.search(r"## Abstract\n(.*?)\n\n\*\*Word count", abstract_md, re.DOTALL)
    if abs_match:
        add_heading(doc, "Abstract", 2)
        abs_text = clean_inline_math(abs_match.group(1).strip())
        add_paragraph(doc, abs_text, space_after_pt=6)
        add_paragraph(doc,
                      "Keywords: agricultural sentiment, VUCA, uncertainty measurement, "
                      "text analysis, Ag Economy Barometer, Granger causality",
                      italic=True)
        add_paragraph(doc, "JEL codes: Q12, Q14, D81, C38", italic=True)
    else:
        parse_and_add(doc, abstract_md)

    doc.add_page_break()

    # ── Sections ──────────────────────────────────────────────────
    sections = [
        ("introduction_draft.md",      "1. Introduction"),
        ("literature_review_draft.md", "2. Background and Related Literature"),
        ("methods_section_draft.md",   "3. Data and Methods"),
        ("results_section_draft.md",   "4. Results"),
        ("discussion_draft.md",        "5. Discussion and Limitations"),
    ]

    for filename, sec_title in sections:
        print(f"  Adding: {sec_title}")
        md = read_draft(filename)
        # Strip the file-level H1 title (it duplicates sec_title)
        md = re.sub(r"^# .+\n", "", md, count=1)
        parse_and_add(doc, md, skip_draft_header=True, skip_notes=True)
        doc.add_paragraph()  # breathing room

    # ── Conclusion ────────────────────────────────────────────────
    print("  Adding: 6. Conclusion")
    conc_md = read_draft("conclusion_abstract_draft.md")
    # Extract just the conclusion section
    conc_match = re.search(r"(## 6\. Conclusion.*?)(?=\n---|\Z)", conc_md, re.DOTALL)
    if conc_match:
        parse_and_add(doc, conc_match.group(1), skip_draft_header=True, skip_notes=True)
    doc.add_paragraph()

    # ── References ────────────────────────────────────────────────
    print("  Adding: References")
    doc.add_page_break()
    add_heading(doc, "References", 2)
    for ref_line in REFERENCES.strip().splitlines():
        if ref_line.strip():
            add_paragraph(doc, ref_line.strip(), space_after_pt=4)
        else:
            doc.add_paragraph()

    # ── Tables ────────────────────────────────────────────────────
    print("  Adding: Tables")
    doc.add_page_break()
    add_heading(doc, "Tables", 2)
    add_paragraph(doc, "[Tables follow — formatted from table1_measures_summary.md]",
                  italic=True, space_after_pt=6)
    tables_md = read_draft("table1_measures_summary.md")
    parse_and_add(doc, tables_md, skip_draft_header=True, skip_notes=True)

    # ── Figures note ──────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Figures", 2)
    fig_notes = [
        "Figure 1. Monthly Ag Economy Barometer index, October 2015–December 2024, "
        "with event annotations. [output/figures/fig1_aeb_vs_time.pdf]",
        "Figure 2. VUCA dimension scores over time (2×2 panel). "
        "[output/figures/fig2_vuca_panel.pdf]",
        "Figure 3. Quantitative vs. text-based sub-measure overlays by dimension. "
        "[output/figures/fig3_subcomponents.pdf]",
        "Figure 4. Cross-correlation heatmap of VUCA dimensions and AEB. "
        "[output/figures/fig4_correlation_heatmap.pdf]",
        "Figure 5. Robustness: 12-month vs. 6-month rolling window for Volatility. "
        "[output/figures/fig5_robustness.pdf]",
        "Figure 6. Impulse-response functions: (a) Volatility -> Farmland Prices, "
        "(b) Ambiguity -> Loan Volume, (c) Uncertainty -> Farmland Prices. "
        "[output/figures/fig6_v_farmland.pdf, fig6_a_loans.pdf, fig6_u_farmland.pdf]",
    ]
    for note in fig_notes:
        add_paragraph(doc, note, space_after_pt=6)

    # ── Save ──────────────────────────────────────────────────────
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"\n[SAVED] {OUTPUT_DOCX}")
    print("[DONE] convert_to_word.py complete.")
    print()
    print("Next steps before submission:")
    print("  1. Render equations: in Word, click after each LaTeX line,")
    print("     press Alt+=, paste the LaTeX source, press Enter")
    print("  2. Resolve remaining [CHECK] items (see notes in each section)")
    print("  3. Insert actual figure image files into the Figures section")
    print("  4. Verify reference list against journal style guide")
    print("  5. Check AJAE word/page limits")


if __name__ == "__main__":
    main()
